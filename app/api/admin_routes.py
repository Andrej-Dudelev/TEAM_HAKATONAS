from __future__ import annotations
import os
import io
import json
from typing import List

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
import pandas as pd
import docx
from pypdf import PdfReader

from app.db.session import get_db
from app.db.models.qa import QAPair, QuestionVariation
from app.db.models.documents import Document
from app.db.models.training import TrainingCourse, TrainingSection, Lesson
from app.schemas.qa import QACreate, QAUpdate, QAOut, QAListOut, VariationIn
from app.schemas.document import DocumentOut, DocumentListOut
from app.schemas.training import CourseCreate, SectionCreate, LessonCreate, LessonUpdate, CourseOut
from app.services.semantic_search import get_service

router = APIRouter(tags=["Admin"])

UPLOAD_DIR = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

def _ensure_semantic():
    svc = get_service()
    if svc is None:
        raise HTTPException(status_code=400, detail="Semantic search service is not available.")
    return svc

def _chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
    text = " ".join(text.split())
    chunks: List[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_size, n)
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        if end == n:
            break
        start = max(0, end - overlap)
    return chunks

def _read_pdf_text(raw: bytes) -> str:
    text = ""
    try:
        reader = PdfReader(io.BytesIO(raw))
        parts: List[str] = [p.extract_text() or "" for p in reader.pages]
        text = "\n".join(parts)
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def _read_docx_text(file_content: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_content))
        return "\n".join([para.text for para in doc.paragraphs if para.text])
    except Exception as e:
        print(f"Error reading docx: {e}")
        return ""

@router.get("/qa", response_model=QAListOut, summary="Get all Q&A pairs with pagination")
def get_all_qa_pairs(skip: int = 0, limit: int = 10, db: Session = Depends(get_db)):
    query = db.query(QAPair).order_by(QAPair.created_at.desc())
    total = query.count()
    items = query.offset(skip).limit(limit).all()
    return {"total": total, "items": items}

@router.post("/qa", response_model=QAOut, summary="Create a Q&A pair")
def create_qa(payload: QACreate, db: Session = Depends(get_db)):
    qa = QAPair(question=payload.question, answer=payload.answer)
    db.add(qa)
    db.flush()
    
    variations_to_add = payload.variations
    if payload.question and not any(v.variation_text == payload.question for v in variations_to_add):
        variations_to_add.append(VariationIn(variation_text=payload.question, language='lt'))

    for v in variations_to_add:
        db.add(QuestionVariation(qa_pair_id=qa.qa_id, variation_text=v.variation_text, language=(v.language or "lt").lower()))
    
    db.commit()
    db.refresh(qa)

    if payload.index:
        svc = _ensure_semantic()
        svc.add_qa_pair(qa)
        
    return qa

@router.post("/qa/upload-file", summary="Upload Q&A pairs from a file")
async def upload_qa_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    svc = _ensure_semantic()
    filename = file.filename or ""
    if not filename.endswith(('.csv', '.xlsx')):
        raise HTTPException(status_code=400, detail="Invalid file type. Please upload a CSV or XLSX file.")

    try:
        content = await file.read()
        if filename.endswith('.xlsx'):
            df = pd.read_excel(io.BytesIO(content))
        else:
            df = pd.read_csv(io.StringIO(content.decode('utf-8')))
        
        df.columns = [str(col).lower() for col in df.columns]

        if 'question' not in df.columns or 'answer' not in df.columns:
            raise HTTPException(status_code=400, detail="File must contain 'question' and 'answer' columns.")

        new_pairs_count = 0
        pairs_to_index = []
        for _, row in df.iterrows():
            question_text = row.get('question')
            answer_text = row.get('answer')
            
            if pd.isna(question_text) or pd.isna(answer_text) or not str(question_text).strip() or not str(answer_text).strip():
                continue

            qa = QAPair(question=str(question_text), answer=str(answer_text))
            db.add(qa)
            db.flush() 
            
            variation = QuestionVariation(
                qa_pair_id=qa.qa_id,
                variation_text=str(question_text),
                language='lt'
            )
            db.add(variation)
            db.flush()
            db.refresh(qa)
            pairs_to_index.append(qa)
            new_pairs_count += 1
        
        db.commit()

        for qa_pair in pairs_to_index:
             svc.add_qa_pair(qa_pair)

        return {"message": f"Successfully imported {new_pairs_count} Q&A pairs."}

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")


@router.get("/docs", response_model=DocumentListOut, summary="Get all indexed documents")
def get_all_documents(db: Session = Depends(get_db)):
    docs = db.query(Document).order_by(Document.created_at.desc()).all()
    return {"items": docs}


@router.post("/docs/upload", response_model=DocumentOut, summary="Upload and index a PDF or DOCX document")
async def upload_document(language: str = Form(..., min_length=2, max_length=2), file: UploadFile = File(...), db: Session = Depends(get_db)):
    name = file.filename or ""
    ext = os.path.splitext(name)[1].lower()
    
    if ext not in [".pdf", ".docx"]:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are allowed.")
    
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file uploaded.")
    
    doc = Document(filename=name, language=language, status="uploading")
    db.add(doc)
    db.commit()
    db.refresh(doc)
    
    text = ""
    if ext == ".pdf":
        text = _read_pdf_text(data)
    elif ext == ".docx":
        text = _read_docx_text(data)

    if not text.strip():
        doc.status = "error_no_text"
        db.commit()
        raise HTTPException(status_code=400, detail=f"Could not extract text from {ext.upper()} file.")
    
    chunks = _chunk_text(text)
    svc = _ensure_semantic()
    svc.index_document_chunks(chunks, document_id=doc.id, language=language)
    
    doc.status = "indexed"
    db.commit()
    db.refresh(doc)
    
    return doc

@router.post("/training/courses", response_model=CourseOut)
def create_course(course: CourseCreate, db: Session = Depends(get_db)):
    db_course = TrainingCourse(
        title=course.title, 
        description=course.description,
        compiler=course.compiler
    )
    db.add(db_course)
    db.commit()
    db.refresh(db_course)
    return db_course

@router.post("/training/sections")
def create_section(section: SectionCreate, db: Session = Depends(get_db)):
    db_section = TrainingSection(course_id=section.course_id, title=section.title, order=section.order)
    db.add(db_section)
    db.commit()
    db.refresh(db_section)
    return db_section

@router.post("/training/lessons")
def create_lesson(lesson: LessonCreate, db: Session = Depends(get_db)):
    criteria = lesson.validation_criteria
    if criteria and "expected_output" in criteria:
        criteria["expected_output"] = criteria["expected_output"].strip()

    db_lesson = Lesson(
        section_id=lesson.section_id, 
        title=lesson.title, 
        content=lesson.content, 
        starting_code=lesson.starting_code,
        order=lesson.order, 
        validation_criteria=criteria
    )
    db.add(db_lesson)
    db.commit()
    db.refresh(db_lesson)
    return db_lesson

@router.put("/training/lessons/{lesson_id}")
def update_lesson(lesson_id: str, payload: LessonUpdate, db: Session = Depends(get_db)):
    db_lesson = db.query(Lesson).filter(Lesson.id == lesson_id).first()
    if not db_lesson:
        raise HTTPException(status_code=404, detail="Lesson not found")

    update_data = payload.dict(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_lesson, key, value)
    
    db.commit()
    db.refresh(db_lesson)
    return db_lesson

@router.post("/training/sections/{section_id}/upload-lessons", summary="Upload lessons from a file")
async def upload_lessons_for_section(section_id: str, file: UploadFile = File(...), db: Session = Depends(get_db)):
    section = db.query(TrainingSection).filter(TrainingSection.id == section_id).first()
    if not section:
        raise HTTPException(status_code=404, detail="Section not found")

    filename = file.filename or ""
    if not filename.endswith(('.csv', '.xlsx')):
        raise HTTPException(status_code=400, detail="Invalid file type. Please upload a CSV or XLSX file.")

    try:
        content = await file.read()
        if filename.endswith('.xlsx'):
            df = pd.read_excel(io.BytesIO(content))
        else:
            df = pd.read_csv(io.StringIO(content.decode('utf-8')))

        df = df.where(pd.notna(df), None)
        df.columns = [str(col).lower().strip() for col in df.columns]
        required_columns = ['title', 'content']
        if not all(col in df.columns for col in required_columns):
            raise HTTPException(status_code=400, detail=f"File must contain the following columns: {', '.join(required_columns)}")

        new_lessons_count = 0
        for _, row in df.iterrows():
            if not row.get('title') or not row.get('content'):
                continue

            criteria = None
            criteria_str = row.get('validation_criteria')
            if criteria_str and isinstance(criteria_str, str):
                try:
                    criteria = json.loads(criteria_str)
                except json.JSONDecodeError:
                    print(f"Skipping malformed JSON in validation_criteria for title: {row.get('title')}")
            
            order_val = row.get('order')
            try:
                order = int(order_val)
            except (ValueError, TypeError):
                order = 0

            lesson = Lesson(
                section_id=section_id,
                title=str(row.get('title')),
                content=str(row.get('content')),
                starting_code=str(row.get('starting_code')) if row.get('starting_code') else None,
                order=order,
                validation_criteria=criteria
            )
            db.add(lesson)
            new_lessons_count += 1
        
        db.commit()
        return {"message": f"Successfully imported {new_lessons_count} lessons into section '{section.title}'."}

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")