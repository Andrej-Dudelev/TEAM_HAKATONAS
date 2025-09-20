
from __future__ import annotations
import os
import io
from typing import List, Optional, Tuple

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from pydantic import BaseModel, Field
from pydantic import ConfigDict  # pydantic v2
from sqlalchemy.orm import Session

from app.db.session import get_db
from app.db.models.qa import QAPair, QuestionVariation
from app.db.models.documents import Document  # jei pas tave yra 'document.py', pakeisk į: from app.db.models.document import Document

# Semantinės paieškos servisas (lazy)
try:
    from app.services.semantic_search import get_service
except Exception:
    get_service = None

router = APIRouter(tags=["QA & Documents"])

# ----------------------- Pydantic schemos (LT) -----------------------

class VariacijaIn(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    variacijos_tekstas: str = Field(..., min_length=1, description="Variacijos tekstas")
    kalba: Optional[str] = Field("lt", min_length=2, max_length=2, description="Kalbos kodas, pvz. 'lt'")

class QAPoraSukurti(BaseModel):
    """
    API priima LT laukus:
      - klausimas (nebūtinas)
      - atsakymas (privalomas)
      - variacijos: [{ variacijos_tekstas, kalba?='lt' }]
      - indeksuoti (bool)
    Viduje įrašoma į abudu DB laukus (*_en ir *_ka), kad nereikėtų migruoti schemos.
    """
    model_config = ConfigDict(populate_by_name=True)

    klausimas: Optional[str] = Field(None, description="Klausimas (nebūtinas)")
    atsakymas: str = Field(..., description="Atsakymas (privalomas)")
    variacijos: List[VariacijaIn] = Field(default_factory=list, description="Papildomos klausimo formuluotės")
    indeksuoti: bool = Field(True, description="Ar iškart indeksuoti į semantinę paiešką")

class QAPoraAtnaujinti(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    klausimas: Optional[str] = Field(None)
    atsakymas: Optional[str] = Field(None)
    variacijos: Optional[List[VariacijaIn]] = Field(None)
    perindeksuoti: bool = Field(True)

class QAPoraOut(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    qa_id: str = Field(..., alias="qa_id")
    klausimas: Optional[str] = None
    atsakymas: str

class DokumentasOut(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    id: str = Field(..., alias="id")
    failo_pavadinimas: str = Field(..., alias="failo_pavadinimas")
    kalba: str = Field(..., alias="kalba")
    busena: str = Field(..., alias="busena")

class DokumentuSarasasOut(BaseModel):
    model_config = ConfigDict(populate_by_name=True)
    irasai: List[DokumentasOut]

# -------------------------- Helpers ----------------------------

UPLOAD_DIR = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

def _ensure_semantic():
    if get_service is None:
        raise HTTPException(status_code=400, detail="Semantic search service is not available.")
    return get_service()

def _chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
    text = " ".join(text.split())
    chunks: List[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_size, n)
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        if end == n:
            break
        start = max(0, end - overlap)
    return chunks

def _pick(*vals: Optional[str]) -> Optional[str]:
    """Paimk pirmą ne-tuščią reikšmę."""
    for v in vals:
        if v and str(v).strip():
            return str(v).strip()
    return None

# -------------------- QA: CRUD (tik LT) --------------------

@router.post("/qa", response_model=QAPoraOut, summary="Sukurti Q&A porą (LT: klausimas/atsakymas)")
def create_qa(payload: QAPoraSukurti, db: Session = Depends(get_db)):
    # LT tekstą įrašom į abu DB laukus, kad atitiktų NOT NULL (answer_en/answer_ka)
    klausimas = payload.klausimas
    atsakymas = payload.atsakymas

    qa = QAPair(
        question_lt=klausimas,
        answer_lt=atsakymas,
      
    )
    db.add(qa)
    db.flush()  # turėti qa_id

    for v in payload.variacijos:
        db.add(QuestionVariation(
            qa_pair_id=qa.qa_id,
            variation_text=v.variacijos_tekstas,
            language=(v.kalba or "lt").lower()
        ))

    db.commit()
    db.refresh(qa)

    if payload.indeksuoti:
        svc = _ensure_semantic()
        svc.add_qa_pair(qa)

    return QAPoraOut(
        qa_id=qa.qa_id,
        klausimas=_pick(qa.question_lt),
        atsakymas=_pick(qa.answer_lt) or "",
    )

@router.get("/qa", response_model=List[QAPoraOut], summary="Gauti Q&A sąrašą (LT laukai)")
def list_qa(skip: int = 0, limit: int = 50, db: Session = Depends(get_db)):
    items = (
        db.query(QAPair)
        .order_by(QAPair.created_at.desc())
        .offset(skip)
        .limit(limit)
        .all()
    )
    out: List[QAPoraOut] = []
    for i in items:
        out.append(QAPoraOut(
            qa_id=i.qa_id,
            klausimas=_pick(i.question_lt),
            atsakymas=_pick(i.answer_lt) or "",
        ))
    return out

@router.get("/qa/{qa_id}", response_model=QAPoraOut, summary="Gauti Q&A pagal ID (LT laukai)")
def get_qa(qa_id: str, db: Session = Depends(get_db)):
    qa = db.query(QAPair).filter(QAPair.qa_id == qa_id).first()
    if not qa:
        raise HTTPException(status_code=404, detail="QAPair not found")
    return QAPoraOut(
        qa_id=qa.qa_id,
        klausimas=_pick(qa.question_lt),
        atsakymas=_pick(qa.answer_lt) or "",
    )

@router.put("/qa/{qa_id}", response_model=QAPoraOut, summary="Atnaujinti Q&A (LT laukai)")
def update_qa(qa_id: str, payload: QAPoraAtnaujinti, db: Session = Depends(get_db)):
    qa = db.query(QAPair).filter(QAPair.qa_id == qa_id).first()
    if not qa:
        raise HTTPException(status_code=404, detail="QAPair not found")

    if payload.klausimas is not None:
        qa.question_lt = payload.klausimas

    if payload.atsakymas is not None:
        qa.answer_lt = payload.atsakymas


    if payload.variacijos is not None:
        db.query(QuestionVariation).filter(QuestionVariation.qa_pair_id == qa_id).delete()
        for v in payload.variacijos:
            db.add(QuestionVariation(
                qa_pair_id=qa_id,
                variation_text=v.variacijos_tekstas,
                language=(v.kalba or "lt").lower()
            ))

    db.commit()
    db.refresh(qa)

    if payload.perindeksuoti:
        svc = _ensure_semantic()
        svc.update_qa_pair(qa)

    return QAPoraOut(
        qa_id=qa.qa_id,
        klausimas=_pick(qa.question_lt),
        atsakymas=_pick(qa.answer_lt) or "",
    )

@router.delete("/qa/{qa_id}", summary="Ištrinti Q&A")
def delete_qa(qa_id: str, db: Session = Depends(get_db)):
    qa = db.query(QAPair).filter(QAPair.qa_id == qa_id).first()
    if not qa:
        raise HTTPException(status_code=404, detail="QAPair not found")
    if get_service is not None:
        svc = get_service()
        svc.delete_qa_pair(qa_id)
    db.delete(qa)
    db.commit()
    return {"status": "ok"}

# -------------------- QA: Įkėlimas iš Excel/Word (LT stulpeliai) --------------------

@router.post(
    "/qa/upload",
    summary="Įkelti Q&A iš Excel (.xlsx/.xls) arba Word (.docx) su LT laukais ('klausimas', 'atsakymas')",
)
async def upload_qa_file(
    failas: UploadFile = File(..., description="Excel (.xlsx/.xls) arba Word (.docx)"),
    indeksuoti: bool = Form(True, description="Ar iškart indeksuoti į semantinę paiešką"),
    db: Session = Depends(get_db),
):
    name = failas.filename or ""
    ext = os.path.splitext(name)[1].lower()

    if ext not in {".xlsx", ".xls", ".docx"}:
        raise HTTPException(status_code=400, detail="Leidžiami: .xlsx, .xls, .docx")

    data = await failas.read()
    created, errors = 0, []

    if ext in {".xlsx", ".xls"}:
        created, errors = _import_qa_from_excel_lt(data, db, do_index=indeksuoti)
    elif ext == ".docx":
        created, errors = _import_qa_from_docx_lt(data, db, do_index=indeksuoti)

    return {"status": "ok", "importuota": created, "klaidos": errors}

def _import_qa_from_excel_lt(raw: bytes, db: Session, do_index: bool) -> Tuple[int, List[str]]:
    import pandas as pd
    try:
        xls = pd.ExcelFile(io.BytesIO(raw))
    except Exception as e:
        return 0, [f"Excel skaitymo klaida: {e}"]

    created = 0
    errors: List[str] = []

    # Sheet 'qa' (arba pirmas) su stulpeliais: 'klausimas' (nebūtina), 'atsakymas' (privaloma)
    sheet_name = "qa" if "qa" in [s.lower() for s in xls.sheet_names] else xls.sheet_names[0]
    df = pd.read_excel(xls, sheet_name=sheet_name)
    cols = {c.lower(): c for c in df.columns}

    if "atsakymas" not in cols:
        return 0, [f"Trūksta stulpelio 'atsakymas' (sheet '{sheet_name}')"]

    for idx, row in df.iterrows():
        klausimas = str(row[cols["klausimas"]]).strip() if "klausimas" in cols and not pd.isna(row.get(cols["klausimas"])) else None
        atsakymas = str(row[cols["atsakymas"]]).strip() if not pd.isna(row.get(cols["atsakymas"])) else ""

        if not atsakymas:
            errors.append(f"Eilutė {idx+2}: privalomas 'atsakymas'. Praleista.")
            continue

        qa = QAPair(
            question_lt=klausimas,
      
            answer_lt=atsakymas,
         
        )
        db.add(qa)
        db.flush()
        created += 1

        if do_index and get_service is not None:
            svc = get_service()
            svc.add_qa_pair(qa)

    db.commit()

    # (nebūtina) – jei norėsi LT variacijų atskiroje skiltyje,
    # gali pridėti antrą sheet, pvz. 'variacijos' (qa_id, variacijos_tekstas, kalba)
    if "variacijos" in [s.lower() for s in xls.sheet_names]:
        vdf = pd.read_excel(xls, sheet_name=[s for s in xls.sheet_names if s.lower() == "variacijos"][0])
        vcols = {c.lower(): c for c in vdf.columns}
        vrequired = ["qa_id", "variacijos_tekstas"]
        vmissing = [c for c in vrequired if c not in vcols]
        if vmissing:
            errors.append(f"Sheet 'variacijos' praleistas: trūksta {', '.join(vmissing)}")
        else:
            count_var = 0
            for idx, row in vdf.iterrows():
                qa_id = row.get(vcols["qa_id"])
                tekstas = row.get(vcols["variacijos_tekstas"])
                kalba = (row.get(vcols["kalba"]) if "kalba" in vcols else "lt")
                if qa_id is None or tekstas is None:
                    errors.append(f"Variacijos eilutė {idx+2}: trūksta 'qa_id' arba 'variacijos_tekstas'. Praleista.")
                    continue
                db.add(QuestionVariation(
                    qa_pair_id=str(qa_id).strip(),
                    language=str(kalba).strip().lower() if kalba else "lt",
                    variation_text=str(tekstas).strip()
                ))
                count_var += 1
            db.commit()
            if do_index and get_service is not None:
                svc = get_service()
                unique_ids = sorted(set([str(v).strip() for v in vdf[vcols["qa_id"]].tolist() if pd.notna(v)]))
                for qid in unique_ids:
                    qa = db.query(QAPair).filter(QAPair.qa_id == qid).first()
                    if qa:
                        svc.update_qa_pair(qa)
                errors.append(f"Variacijų įkelta: {count_var} (perindeksuota {len(unique_ids)} Q&A).")

    return created, errors

def _import_qa_from_docx_lt(raw: bytes, db: Session, do_index: bool) -> Tuple[int, List[str]]:
    """
    Tikimės paprasto formato DOCX (LT):
      Klausimas: ...
      Atsakymas: ...
    arba trumpiau:
      K: ...
      A: ...
    Poros kartojasi dokumente.
    """
    from docx import Document as DocxDocument

    errors: List[str] = []
    created = 0
    doc = DocxDocument(io.BytesIO(raw))

    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    def _norm(s: str) -> str:
        return s.strip().lstrip("-").strip()

    i = 0
    while i < len(paras):
        q = a = None
        low = paras[i].lower()
        if low.startswith(("klausimas:", "k:", "q:", "q (lt):")):
            q = _norm(paras[i].split(":", 1)[1]); i += 1
        if i < len(paras):
            low2 = paras[i].lower()
            if low2.startswith(("atsakymas:", "a:", "ans:", "a (lt):")):
                a = _norm(paras[i].split(":", 1)[1]); i += 1

        if a:
            qa = QAPair(
                question_lt=q,  # saugom abu
              
                answer_lt=a,
             
            )
            db.add(qa)
            db.flush()
            created += 1
            if do_index and get_service is not None:
                svc = get_service()
                svc.add_qa_pair(qa)
        else:
            # jei neatitiko maskės – praleidžiam šį paragrafą
            i += 1

    db.commit()
    return created, errors

# ---------------- Dokumentai: PDF įkėlimas ir indeksavimas ----------------

@router.post("/docs/upload", response_model=DokumentasOut, summary="Įkelti PDF dokumentą ir indeksuoti")
async def upload_document_pdf(
    kalba: str = Form(..., min_length=2, max_length=2, description="Kalbos kodas, pvz. 'lt'"),
    failas: UploadFile = File(..., description="PDF dokumentas"),
    db: Session = Depends(get_db),
):
    name = failas.filename or ""
    ext = os.path.splitext(name)[1].lower()
    if ext != ".pdf":
        raise HTTPException(status_code=400, detail="Leidžiamas tik PDF.")

    data = await failas.read()
    if not data:
        raise HTTPException(status_code=400, detail="Tuščias failas.")

    target_path = os.path.join(UPLOAD_DIR, name)
    try:
        with open(target_path, "wb") as f:
            f.write(data)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failo išsaugojimo klaida: {e}")

    doc = Document(filename=name, language=kalba, status="uploaded")
    db.add(doc)
    db.commit()
    db.refresh(doc)

    text = _read_pdf_text(data)
    if not text.strip():
        raise HTTPException(status_code=400, detail="Nepavyko ištraukti teksto iš PDF.")

    chunks = _chunk_text(text)
    if get_service is None:
        raise HTTPException(status_code=400, detail="Semantic search service is not available.")
    svc = get_service()
    svc.index_document_chunks(chunks, document_id=doc.id, language=kalba)

    doc.status = "indexed"
    db.commit()
    db.refresh(doc)

    return DokumentasOut(id=doc.id, failo_pavadinimas=doc.filename, kalba=doc.language, busena=doc.status)

def _read_pdf_text(raw: bytes) -> str:
    """
    Naudoja pypdf (PdfReader). Jei nepavyksta, bando pdfminer.six.
    """
    text = ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        parts: List[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
        text = "\n".join(parts)
    except Exception:
        try:
            from pdfminer.high_level import extract_text
            text = extract_text(io.BytesIO(raw))
        except Exception:
            text = ""
    return text

# ---------------- Dokumentų sąrašas / get / delete / reindex ----------------

@router.get("/docs", response_model=DokumentuSarasasOut, summary="Sąrašas dokumentų")
def list_documents(db: Session = Depends(get_db)):
    items = db.query(Document).order_by(Document.created_at.desc()).all()
    return DokumentuSarasasOut(
        irasai=[DokumentasOut(id=i.id, failo_pavadinimas=i.filename, kalba=i.language, busena=i.status) for i in items]
    )

@router.get("/docs/{doc_id}", response_model=DokumentasOut, summary="Gauti dokumentą")
def get_document(doc_id: str, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return DokumentasOut(id=doc.id, failo_pavadinimas=doc.filename, kalba=doc.language, busena=doc.status)

@router.delete("/docs/{doc_id}", summary="Trinti dokumentą (tik DB įrašą)")
def delete_document(doc_id: str, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    db.delete(doc)
    db.commit()
    # (pasirinktinai) – gali trinti embeddings ChromaDB pagal document_id, jei servise taip saugai metadatas
    return {"status": "ok", "deleted": doc_id}

@router.post("/docs/{doc_id}/reindex", summary="Perindeksuoti PDF dokumentą iš naujo")
def reindex_document(doc_id: str, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    path = os.path.join(UPLOAD_DIR, doc.filename)
    if not os.path.exists(path):
        raise HTTPException(status_code=400, detail="Failas nerastas diske – įkelk iš naujo.")

    with open(path, "rb") as f:
        raw = f.read()
    text = _read_pdf_text(raw)
    if not text.strip():
        raise HTTPException(status_code=400, detail="Nepavyko ištraukti teksto iš PDF.")

    chunks = _chunk_text(text)
    if get_service is None:
        raise HTTPException(status_code=400, detail="Semantic search service is not available.")
    svc = get_service()
    svc.index_document_chunks(chunks, document_id=doc.id, language=doc.language)

    doc.status = "indexed"
    db.commit()
    db.refresh(doc)
    return {"status": "ok", "reindexed": doc_id, "chunks": len(chunks)}
